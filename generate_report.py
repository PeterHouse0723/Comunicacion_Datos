from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Styles helpers ────────────────────────────────────────────────────────────
def set_heading(para, text, level=1):
    para.text = text
    para.style = f'Heading {level}'

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return p

def add_para(doc, text, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p

def add_table_row(table, cells_data, is_header=False):
    row = table.add_row()
    for i, val in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = str(val)
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(str(val))
        run.font.size = Pt(9.5)
        if is_header:
            run.bold = True
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '2E74B5')
            shading.set(qn('w:color'), 'FFFFFF')
            shading.set(qn('w:val'), 'clear')
            cell._tc.get_or_add_tcPr().append(shading)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return row

# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('INFORME TÉCNICO DEL PROYECTO')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run('Sistema de Gestión Académica\nComunicacion_Datos')
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = meta.add_run(f'Fecha de generación: {datetime.datetime.now().strftime("%d de %B de %Y")}')
run3.font.size = Pt(11)
run3.italic = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. DESCRIPCIÓN GENERAL
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '1. Descripción General del Sistema', 1)
add_para(doc,
    'El sistema "Comunicacion_Datos" es una plataforma web de gestión académica multi-institución '
    'desarrollada con arquitectura cliente-servidor. Permite administrar instituciones educativas, '
    'sus cursos, docentes, estudiantes, calificaciones, asistencias y comunicación entre actores '
    'académicos. Cuenta con un chatbot integrado para consultas rápidas.')

add_heading(doc, '1.1 Arquitectura General', 2)
add_para(doc, 'El sistema sigue el patrón Cliente-Servidor con renderizado del lado del servidor (SSR):')
add_bullet(doc, 'Cliente: Navegador web — HTML/CSS/JavaScript (Jinja2 templates)')
add_bullet(doc, 'Servidor: Flask (Python) con ORM SQLAlchemy')
add_bullet(doc, 'Base de datos: SQLite (desarrollo) / PostgreSQL (producción)')
add_bullet(doc, 'Despliegue: Render.com con Gunicorn como servidor WSGI')

# ══════════════════════════════════════════════════════════════════════════════
# 2. STACK TECNOLÓGICO
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '2. Stack Tecnológico', 1)

add_heading(doc, '2.1 Frontend', 2)
items_fe = [
    ('Motor de plantillas', 'Jinja2 (Flask built-in)'),
    ('Estilos', 'CSS personalizado (login.css, admin.css, docente.css, estudiante.css)'),
    ('JavaScript', 'Vanilla JS (sin frameworks)'),
    ('Funcionalidades UI', 'Formularios, modales, gráficas, dashboards por rol'),
]
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
add_table_row(t, ['Componente', 'Tecnología'], is_header=True)
for item in items_fe:
    add_table_row(t, item)
doc.add_paragraph()

add_heading(doc, '2.2 Backend', 2)
items_be = [
    ('Framework web', 'Flask 3.0.0'),
    ('ORM', 'SQLAlchemy 2.0.47 + Flask-SQLAlchemy 3.1.1'),
    ('Migraciones', 'Flask-Migrate 4.1.0 (Alembic)'),
    ('Servidor WSGI (prod)', 'Gunicorn 21.2.0'),
    ('Adaptador PostgreSQL', 'psycopg2-binary 2.9.11'),
    ('Variables de entorno', 'python-dotenv 1.0.1'),
    ('Utilidades WSGI', 'Werkzeug 3.1.3'),
    ('CLI commands', 'Flask-Script 2.0.6'),
]
t2 = doc.add_table(rows=1, cols=2)
t2.style = 'Table Grid'
add_table_row(t2, ['Componente', 'Tecnología'], is_header=True)
for item in items_be:
    add_table_row(t2, item)
doc.add_paragraph()

add_heading(doc, '2.3 Base de Datos', 2)
add_bullet(doc, 'Desarrollo: SQLite (archivo local backend/app.db)')
add_bullet(doc, 'Producción: PostgreSQL (Render.com, conexión via DATABASE_URL)')
add_bullet(doc, 'Migraciones SQL manuales: backend/migrations/*.sql (7 archivos)')

# ══════════════════════════════════════════════════════════════════════════════
# 3. ESTRUCTURA DEL PROYECTO
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '3. Estructura del Proyecto', 1)
add_para(doc, 'El repositorio está organizado de la siguiente forma:')

structure = '''
Comunicacion_Datos/
├── backend/
│   ├── app.py                  # Factory de la aplicación Flask
│   ├── config.py               # Configuraciones (Dev/Prod/Test)
│   ├── extensions.py           # Instancias de extensiones (db, migrate)
│   ├── models.py               # Modelos ORM (17 tablas)
│   ├── utils.py                # Funciones auxiliares (auth, validación)
│   ├── wsgi.py                 # Punto de entrada WSGI (producción)
│   ├── manage.py               # Comandos CLI de gestión
│   ├── requirements.txt        # Dependencias Python
│   ├── routes/
│   │   ├── auth.py             # Rutas de autenticación
│   │   ├── dashboard.py        # Rutas por rol (admin/docente/estudiante)
│   │   └── admin.py            # Rutas panel de administración
│   ├── templates/              # HTML Jinja2 por rol
│   │   ├── login.html / register.html
│   │   ├── dashboard/
│   │   ├── admin/ / admin_local/
│   │   ├── docente/
│   │   ├── estudiante/
│   │   └── components/chatbot.html
│   ├── static/
│   │   ├── css/                # Hojas de estilo por rol
│   │   ├── js/admin.js
│   │   └── imagenes/
│   ├── chatbot/
│   │   ├── routes.py           # Endpoint POST /api/chatbot/mensaje
│   │   └── engine.py           # Motor NLP por intenciones
│   ├── scripts/
│   │   ├── cargar_estudiantes_cli.py
│   │   └── simular_actividades.py
│   └── migrations/             # Scripts SQL de migración
└── Procfile                    # Configuración Render.com
'''
p = doc.add_paragraph()
run = p.add_run(structure)
run.font.name = 'Courier New'
run.font.size = Pt(8.5)

# ══════════════════════════════════════════════════════════════════════════════
# 4. BASE DE DATOS – ESQUEMA Y MER
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '4. Base de Datos', 1)
add_para(doc,
    'El sistema utiliza 17 tablas relacionales gestionadas mediante SQLAlchemy ORM. '
    'A continuación se describe cada entidad con sus atributos y relaciones principales.')

entities = [
    ('Institucion', 'id, nombre, ciudad, pais, logo_filename, admin_global_id (FK→Usuario), activo, fecha_creacion',
     'Agrupa usuarios, periodos y cursos de una institución.'),
    ('Usuario', 'id, institucion_id (FK), email (unique), password (hash), nombre, apellido, role [admin_global|admin_local|docente|estudiante], estado [pendiente|activo|inactivo|suspendido], contraseña_cambiada, fecha_creacion, fecha_actualizacion',
     'Entidad central de identidad. Todos los roles comparten esta tabla (herencia de tabla única).'),
    ('Periodo', 'id, institucion_id (FK), nombre, fecha_inicio, fecha_fin, activo',
     'Semestre o período académico de una institución.'),
    ('Curso', 'id, institucion_id (FK), periodo_id (FK), nombre, codigo, descripcion, creditos, docente_principal_id (FK→Usuario), activo, dias_semana (JSON), sesiones_por_semana',
     'Materia o asignatura. Genera automáticamente clases según el calendario.'),
    ('CursoDocente', 'id, curso_id (FK), docente_id (FK) — UNIQUE(curso_id, docente_id)',
     'Tabla pivot M:N entre Curso y Usuario (docente).'),
    ('EstudianteCurso', 'id, estudiante_id (FK), curso_id (FK), fecha_inscripcion — UNIQUE(estudiante_id, curso_id)',
     'Tabla pivot M:N entre Usuario (estudiante) y Curso.'),
    ('SolicitudEstudianteMateria', 'id, curso_id (FK), estudiante_id (FK), docente_id (FK), admin_local_id (FK), estado [pendiente|aprobado|rechazado], motivo, respuesta, fecha_solicitud, fecha_resolucion',
     'Solicitud del docente para inscribir un estudiante existente en un curso.'),
    ('SolicitudNuevoEstudiante', 'id, curso_id (FK), docente_id (FK), admin_local_id (FK), nombre, apellido, correo, estado, motivo_rechazo, estudiante_id (FK, post-aprobación), fecha_solicitud, fecha_resolucion',
     'Solicitud del docente para crear e inscribir un nuevo estudiante.'),
    ('Clase', 'id, curso_id (FK), periodo_id (FK), fecha, numero_clase, tema, fecha_creacion',
     'Sesión de clase individual generada a partir del calendario del curso.'),
    ('Nota', 'id, estudiante_id (FK), curso_id (FK), valor_nota (0.0-5.0), numero_entrega, tipo_evaluacion, descripcion, fecha_registro',
     'Modelo legado de calificaciones simples (escala 0-5, aprobación ≥ 3.2).'),
    ('Actividad', 'id, curso_id (FK), nombre, descripcion, tipo_evaluacion, semana, ponderacion, fecha_asignacion, fecha_vencimiento, activa',
     'Evaluación o tarea asociada a un curso con ponderación y fechas.'),
    ('Calificacion', 'id, actividad_id (FK), estudiante_id (FK), valor_nota, retroalimentacion, fecha_calificacion — UNIQUE(actividad_id, estudiante_id)',
     'Nota de un estudiante en una Actividad específica.'),
    ('Asistencia', 'id, estudiante_id (FK), clase_id (FK), curso_id (FK), presente (bool), justificacion, fecha_registro — UNIQUE(estudiante_id, clase_id)',
     'Registro de asistencia por clase y estudiante.'),
    ('AlertaRiesgoAcademico', 'id, estudiante_id (FK), curso_id (FK), tipo_alerta [inasistencia|bajo_promedio|multiples_bajas], promedio_actual, porcentaje_inasistencia, total_notas_bajas, fecha_alerta, estado [activa|resuelta], descripcion',
     'Alerta automática generada al detectar riesgo académico.'),
    ('LoginAuditoria', 'id, usuario_id (FK), fecha_login, ip_address, navegador, estado [exitoso|fallido], razon_fallo',
     'Trazabilidad de intentos de inicio de sesión.'),
    ('Notificacion', 'id, usuario_id (FK), titulo, mensaje, tipo, leida (bool), fecha_creacion, fecha_lectura',
     'Notificaciones internas del sistema para cada usuario.'),
    ('Mensaje', 'id, remitente_id (FK), destinatario_id (FK), curso_id (FK), contenido, leido (bool), fecha_lectura, fecha_creacion',
     'Mensajería bidireccional entre docente y estudiante en el contexto de un curso.'),
]

for name, attrs, desc in entities:
    add_heading(doc, f'4.{entities.index((name,attrs,desc))+1} {name}', 2)
    add_para(doc, desc, italic=True)
    p = doc.add_paragraph()
    run = p.add_run('Atributos: ')
    run.bold = True
    run.font.size = Pt(10)
    p.add_run(attrs).font.size = Pt(10)

doc.add_paragraph()
add_heading(doc, '4.18 Diagrama MER – Representación textual', 2)
add_para(doc, 'Las relaciones principales entre entidades son:')

mer_relations = [
    'Institucion (1) ──── (N) Usuario',
    'Institucion (1) ──── (N) Periodo',
    'Institucion (1) ──── (N) Curso',
    'Periodo    (1) ──── (N) Curso',
    'Periodo    (1) ──── (N) Clase',
    'Curso      (1) ──── (N) Clase',
    'Curso      (M) ──── (N) Usuario[docente]  → via CursoDocente',
    'Curso      (M) ──── (N) Usuario[estudiante] → via EstudianteCurso',
    'Curso      (1) ──── (N) Actividad',
    'Curso      (1) ──── (N) Nota',
    'Curso      (1) ──── (N) Asistencia',
    'Curso      (1) ──── (N) AlertaRiesgoAcademico',
    'Curso      (1) ──── (N) Mensaje',
    'Actividad  (1) ──── (N) Calificacion',
    'Clase      (1) ──── (N) Asistencia',
    'Usuario    (1) ──── (N) Nota',
    'Usuario    (1) ──── (N) Calificacion',
    'Usuario    (1) ──── (N) Asistencia',
    'Usuario    (1) ──── (N) AlertaRiesgoAcademico',
    'Usuario    (1) ──── (N) LoginAuditoria',
    'Usuario    (1) ──── (N) Notificacion',
    'Usuario    (1) ──── (N) Mensaje (remitente / destinatario)',
    'Usuario    (1) ──── (N) SolicitudEstudianteMateria',
    'Usuario    (1) ──── (N) SolicitudNuevoEstudiante',
]
for rel in mer_relations:
    add_bullet(doc, rel)

# ══════════════════════════════════════════════════════════════════════════════
# 5. SEGURIDAD
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '5. Modelo de Seguridad', 1)

add_heading(doc, '5.1 Autenticación', 2)
auth_items = [
    ('Mecanismo', 'Sesiones del lado del servidor (Flask session + cookie firmada)'),
    ('Hash de contraseñas', 'Werkzeug pbkdf2:sha256 (generate_password_hash / check_password_hash)'),
    ('Contraseña temporal', 'Generación aleatoria; el usuario debe cambiarla en el primer acceso'),
    ('Reset de contraseña', 'Token temporal con fecha de expiración (utils.generar_token_reset)'),
    ('Timeout de sesión', '7 días (PERMANENT_SESSION_LIFETIME)'),
]
t3 = doc.add_table(rows=1, cols=2)
t3.style = 'Table Grid'
add_table_row(t3, ['Aspecto', 'Implementación'], is_header=True)
for item in auth_items:
    add_table_row(t3, item)
doc.add_paragraph()

add_heading(doc, '5.2 Validación de Contraseñas', 2)
add_para(doc, 'La función validar_contraseña() en utils.py aplica las siguientes reglas:')
add_bullet(doc, 'Mínimo 8 caracteres')
add_bullet(doc, 'Al menos 1 letra mayúscula')
add_bullet(doc, 'Al menos 1 número')
add_bullet(doc, 'Al menos 1 carácter especial: !@#$%^&*')

add_heading(doc, '5.3 Autorización (Control de Acceso por Roles)', 2)
add_para(doc, 'El sistema implementa RBAC (Role-Based Access Control) con 4 roles:')
roles_data = [
    ('admin_global', 'Gestión total: instituciones, usuarios, cursos, solicitudes'),
    ('admin_local', 'Gestión de su institución: cursos, docentes, solicitudes de estudiantes'),
    ('docente', 'Gestión de sus cursos: asistencia, calificaciones, mensajería, solicitudes'),
    ('estudiante', 'Consulta de sus cursos, calificaciones, asistencia y alertas'),
]
t4 = doc.add_table(rows=1, cols=2)
t4.style = 'Table Grid'
add_table_row(t4, ['Rol', 'Permisos'], is_header=True)
for item in roles_data:
    add_table_row(t4, item)
doc.add_paragraph()
add_para(doc,
    'Los decoradores @login_required y @admin_required verifican session["usuario_id"] y '
    'session["role"] en cada request protegido.')

add_heading(doc, '5.4 Seguridad en Cookies y Sesión', 2)
cookie_items = [
    ('SESSION_COOKIE_HTTPONLY', 'True — previene acceso desde JavaScript'),
    ('SESSION_COOKIE_SAMESITE', '"Lax" — protección CSRF básica'),
    ('SESSION_COOKIE_SECURE', 'True en producción — solo HTTPS'),
    ('SECRET_KEY', 'Clave secreta larga, diferente por entorno'),
]
t5 = doc.add_table(rows=1, cols=2)
t5.style = 'Table Grid'
add_table_row(t5, ['Parámetro', 'Valor / Propósito'], is_header=True)
for item in cookie_items:
    add_table_row(t5, item)
doc.add_paragraph()

add_heading(doc, '5.5 Auditoría de Acceso', 2)
add_para(doc,
    'Cada intento de login (exitoso o fallido) queda registrado en la tabla LoginAuditoria '
    'con: usuario_id, fecha, IP de origen, agente de usuario (navegador), estado y razón de fallo.')

add_heading(doc, '5.6 Estado de Usuarios', 2)
add_para(doc,
    'Los usuarios pueden estar en cuatro estados: pendiente (registro reciente, sin aprobación), '
    'activo (puede iniciar sesión), inactivo o suspendido (acceso bloqueado). '
    'El flujo de aprobación de docentes y admins locales pasa por el administrador.')

# ══════════════════════════════════════════════════════════════════════════════
# 6. ENDPOINTS PRINCIPALES
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '6. Endpoints Principales de la API', 1)

add_heading(doc, '6.1 Autenticación (/auth)', 2)
auth_routes = [
    ('GET/POST', '/auth/login', 'Inicio de sesión'),
    ('GET',      '/auth/logout', 'Cierre de sesión'),
    ('GET/POST', '/auth/register', 'Registro de nuevo usuario'),
    ('GET/POST', '/auth/cambiar-contraseña-obligatorio', 'Cambio obligatorio (primer acceso)'),
    ('GET/POST', '/auth/olvide-contraseña', 'Solicitud de reset de contraseña'),
    ('GET/POST', '/auth/reset-contraseña/<token>', 'Reset con token temporal'),
]
t6 = doc.add_table(rows=1, cols=3)
t6.style = 'Table Grid'
add_table_row(t6, ['Método', 'Ruta', 'Descripción'], is_header=True)
for r in auth_routes:
    add_table_row(t6, r)
doc.add_paragraph()

add_heading(doc, '6.2 Dashboard Docente (/dashboard/docente)', 2)
docente_routes = [
    ('GET',  '/dashboard/docente', 'Panel principal del docente'),
    ('GET',  '/dashboard/docente/cursosdocente', 'Listar cursos del docente'),
    ('GET',  '/dashboard/docente/cursos/<id>', 'Detalle de un curso'),
    ('GET',  '/dashboard/docente/cursos/<id>/calificaciones', 'Calificaciones del curso'),
    ('GET/POST', '/dashboard/docente/cursos/<id>/asistencia', 'Registro de asistencia'),
    ('POST', '/dashboard/docente/solicitudes', 'Solicitar inscripción de estudiante'),
    ('GET/POST', '/dashboard/docente/solicitar-estudiante', 'Solicitar nuevo estudiante'),
    ('GET',  '/dashboard/docente/alertas', 'Ver alertas de riesgo académico'),
]
t7 = doc.add_table(rows=1, cols=3)
t7.style = 'Table Grid'
add_table_row(t7, ['Método', 'Ruta', 'Descripción'], is_header=True)
for r in docente_routes:
    add_table_row(t7, r)
doc.add_paragraph()

add_heading(doc, '6.3 Dashboard Estudiante (/dashboard/estudiante)', 2)
est_routes = [
    ('GET', '/dashboard/estudiante/dashboard', 'Panel principal del estudiante'),
    ('GET', '/dashboard/estudiante/cursos', 'Listar cursos inscritos'),
    ('GET', '/dashboard/estudiante/cursos/<id>/calificaciones', 'Ver calificaciones'),
    ('GET', '/dashboard/estudiante/cursos/<id>/asistencia', 'Ver asistencia'),
    ('GET', '/dashboard/estudiante/alertas', 'Ver alertas académicas'),
]
t8 = doc.add_table(rows=1, cols=3)
t8.style = 'Table Grid'
add_table_row(t8, ['Método', 'Ruta', 'Descripción'], is_header=True)
for r in est_routes:
    add_table_row(t8, r)
doc.add_paragraph()

add_heading(doc, '6.4 Mensajería (/dashboard/api/mensajes)', 2)
msg_routes = [
    ('POST', '/dashboard/api/mensajes', 'Enviar mensaje'),
    ('GET',  '/dashboard/api/mensajes/<usuario_id>/<curso_id>', 'Obtener conversación'),
    ('GET',  '/dashboard/api/mensajes/no-leidos/<curso_id>', 'Contar mensajes no leídos'),
    ('PUT',  '/dashboard/api/mensajes/<id>/marcar-leido', 'Marcar mensaje como leído'),
    ('GET',  '/dashboard/api/mensajes/remitentes/<curso_id>', 'Listar conversaciones activas'),
]
t9 = doc.add_table(rows=1, cols=3)
t9.style = 'Table Grid'
add_table_row(t9, ['Método', 'Ruta', 'Descripción'], is_header=True)
for r in msg_routes:
    add_table_row(t9, r)
doc.add_paragraph()

add_heading(doc, '6.5 Administración (/admin)', 2)
admin_routes = [
    ('GET',  '/admin/dashboard', 'Panel de administración'),
    ('GET',  '/admin/usuarios', 'Gestión de usuarios'),
    ('GET',  '/admin/cursos', 'Gestión de cursos'),
    ('POST', '/admin/cursos/crear', 'Crear nuevo curso'),
    ('POST', '/admin/cursos/<id>/cargar-estudiantes', 'Cargar estudiantes (individual/CSV)'),
    ('GET',  '/admin/solicitudes-estudiantes', 'Ver solicitudes de estudiantes'),
    ('POST', '/admin/solicitudes-estudiantes/<id>/aprobar', 'Aprobar solicitud'),
    ('GET',  '/admin/instituciones', 'Gestión de instituciones'),
    ('POST', '/admin/instituciones/crear', 'Crear institución'),
]
t10 = doc.add_table(rows=1, cols=3)
t10.style = 'Table Grid'
add_table_row(t10, ['Método', 'Ruta', 'Descripción'], is_header=True)
for r in admin_routes:
    add_table_row(t10, r)
doc.add_paragraph()

add_heading(doc, '6.6 Chatbot (/api/chatbot)', 2)
add_bullet(doc, 'POST /api/chatbot/mensaje — Envía un mensaje al chatbot y recibe respuesta')
add_para(doc,
    'El motor NLP (chatbot/engine.py) clasifica las intenciones: ayuda, mis_cursos, '
    'mis_faltas, estado_materia y mensajes, adaptando la respuesta al rol del usuario.')

# ══════════════════════════════════════════════════════════════════════════════
# 7. CONFIGURACIÓN Y ENTORNOS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '7. Configuración y Entornos', 1)
env_data = [
    ('Clase', 'Entorno', 'Base de datos', 'Debug'),
    ('DevelopmentConfig', 'Desarrollo', 'SQLite (backend/app.db)', 'True'),
    ('ProductionConfig', 'Producción', 'PostgreSQL (DATABASE_URL)', 'False'),
    ('TestingConfig', 'Pruebas', 'SQLite en memoria', 'True'),
]
t11 = doc.add_table(rows=1, cols=4)
t11.style = 'Table Grid'
add_table_row(t11, env_data[0], is_header=True)
for row in env_data[1:]:
    add_table_row(t11, row)
doc.add_paragraph()

add_para(doc, 'Variables de entorno requeridas (.env):')
add_bullet(doc, 'DATABASE_URL — Cadena de conexión a la base de datos')
add_bullet(doc, 'SECRET_KEY — Clave secreta para firma de sesiones (mín. 32 chars)')
add_bullet(doc, 'FLASK_ENV — Entorno activo (development / production)')

add_heading(doc, '7.1 Despliegue en Producción (Render.com)', 2)
add_para(doc, 'El archivo Procfile define el comando de arranque:')
p = doc.add_paragraph()
run = p.add_run('web: cd backend && gunicorn wsgi:app --timeout 600 --workers 1 --max-requests 500 --graceful-timeout 120 --keep-alive 65')
run.font.name = 'Courier New'
run.font.size = Pt(9)

# ══════════════════════════════════════════════════════════════════════════════
# 8. FUNCIONALIDADES PRINCIPALES
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '8. Funcionalidades del Sistema', 1)

features = {
    'Gestión Multi-institución': [
        'Creación y administración de instituciones con logo',
        'Períodos académicos por institución',
        'Administrador local por institución',
    ],
    'Gestión de Usuarios': [
        'Registro con validación de contraseña fuerte',
        'Flujo de aprobación para docentes y administradores',
        'Cambio de contraseña obligatorio en primer acceso',
        'Reset de contraseña por token',
        'Estados: pendiente, activo, inactivo, suspendido',
    ],
    'Gestión Académica': [
        'Creación de cursos con horario semanal',
        'Generación automática de clases basada en calendario',
        'Inscripción de estudiantes (individual o por CSV)',
        'Asignación de docentes a cursos',
    ],
    'Calificaciones': [
        'Actividades con ponderación y fechas de vencimiento',
        'Registro de calificaciones por actividad y estudiante',
        'Modelo legado de notas simples (escala 0-5)',
        'Cálculo de promedio por clase',
    ],
    'Asistencia': [
        'Registro de asistencia por fecha de clase',
        'Consulta histórica por curso',
        'Reset de asistencia para correcciones',
    ],
    'Alertas de Riesgo Académico': [
        'Detección automática: inasistencia excesiva, bajo promedio, múltiples notas bajas',
        'Vista para docentes y estudiantes',
        'Estados: activa / resuelta',
    ],
    'Mensajería Interna': [
        'Chat bidireccional docente ↔ estudiante por curso',
        'Indicador de mensajes no leídos',
        'Historial de conversaciones',
    ],
    'Solicitudes de Estudiantes': [
        'Docente solicita inscribir estudiante existente en curso',
        'Docente solicita crear nuevo estudiante',
        'Flujo de aprobación/rechazo por admin local',
    ],
    'Chatbot Integrado': [
        'Consultas sobre cursos, faltas y estado académico',
        'Motor NLP por intenciones (sin ML externo)',
        'Respuestas adaptadas al rol del usuario',
    ],
}

for feat_name, feat_items in features.items():
    add_heading(doc, f'8.x {feat_name}', 2)
    for item in feat_items:
        add_bullet(doc, item)

# ══════════════════════════════════════════════════════════════════════════════
# 9. SCRIPTS Y HERRAMIENTAS AUXILIARES
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '9. Scripts y Herramientas Auxiliares', 1)
add_bullet(doc, 'backend/scripts/cargar_estudiantes_cli.py — Carga masiva de estudiantes por línea de comandos')
add_bullet(doc, 'backend/scripts/simular_actividades.py — Generación de actividades y calificaciones de prueba')
add_bullet(doc, 'backend/manage.py init_db — Inicializa la base de datos')
add_bullet(doc, 'backend/manage.py create_admin — Crea el administrador global por defecto')
add_bullet(doc, 'backend/manage.py migrate_db — Ejecuta migraciones SQL pendientes')

# ══════════════════════════════════════════════════════════════════════════════
# 10. CONCLUSIONES
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '10. Conclusiones', 1)
add_para(doc,
    'El sistema Comunicacion_Datos es una solución web completa de gestión académica que '
    'implementa una arquitectura cliente-servidor con Flask y SQLAlchemy. Ofrece una '
    'separación clara de responsabilidades por rol, un modelo de seguridad robusto '
    '(RBAC, sesiones firmadas, hashing de contraseñas, auditoría de acceso) y una base de '
    'datos normalizada con 17 entidades relacionales.')
add_para(doc,
    'La plataforma está preparada para producción mediante despliegue en Render.com con '
    'PostgreSQL y Gunicorn, contando con múltiples entornos de configuración y scripts de '
    'migración SQL. El chatbot integrado aporta una capa de interacción conversacional '
    'sin dependencias externas de ML.')

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = r'c:\ProCD\Comunicacion_Datos\Informe_Tecnico_Comunicacion_Datos.docx'
doc.save(output_path)
print(f'Documento guardado en: {output_path}')
